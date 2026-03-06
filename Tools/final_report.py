"""
FILE: tools/final_report_tools.py
------------------------------------
LangChain tools exposed to the Final Report ReAct agent.
Generates a comprehensive Word document with ALL pipeline data.
"""

import json
import os
import pandas as pd
from langchain_core.tools import tool

from core.final_report_engine import assemble_report
from Schemas.final_report import FinalReportOutput


# ─────────────────────────────────────────────
# SESSION STORE
# ─────────────────────────────────────────────

_report_store: dict = {
    "original_query":        None,
    "profiler_output":       None,
    "preprocessor_output":   None,
    "methodologist_output":  None,
    "checker_output":        None,
    "statistician_output":   None,
    "rectification_output":  None,
    "critic_output":         None,
    "report_output":         None,
}


def init_report_store(
    original_query: str,
    profiler_output: dict,
    preprocessor_output: dict,
    methodologist_output: dict,
    checker_output: dict,
    statistician_output: dict,
    rectification_output: dict | None = None,
    critic_output: dict | None = None,
) -> None:
    """Called by run_final_report() before invoking the agent."""
    _report_store["original_query"]       = original_query
    _report_store["profiler_output"]      = profiler_output
    _report_store["preprocessor_output"]  = preprocessor_output
    _report_store["methodologist_output"] = methodologist_output
    _report_store["checker_output"]       = checker_output
    _report_store["statistician_output"]  = statistician_output
    _report_store["rectification_output"] = rectification_output
    _report_store["critic_output"]        = critic_output
    _report_store["report_output"]        = None


def get_report_store() -> dict:
    return _report_store


# ─────────────────────────────────────────────
# TOOLS
# ─────────────────────────────────────────────

@tool
def get_pipeline_summary() -> str:
    """
    Returns a structured summary of everything that happened in the pipeline.
    Use this to understand what to cover in the interpretation before writing it.
    """
    store  = _report_store
    stat   = store.get("statistician_output", {})
    meth   = store.get("methodologist_output", {})
    rect   = store.get("rectification_output") or {}
    check  = store.get("checker_output", {})
    crit   = store.get("critic_output") or {}
    prof   = store.get("profiler_output", {})
    family = stat.get("test_family", "")

    result_summary: dict = {}
    if family == "inference" and stat.get("inference_result"):
        r = stat["inference_result"]
        result_summary = {
            "statistic": f"{r.get('statistic_label')}={r.get('statistic')}",
            "p_value": r.get("p_value"), "verdict": r.get("verdict"),
            "effect_size": f"{r.get('effect_size_label')}={r.get('effect_size')}",
            "group_stats": r.get("group_stats", {}),
            "interpretation": r.get("interpretation"),
        }
    elif family == "regression" and stat.get("regression_result"):
        r = stat["regression_result"]
        result_summary = {
            "r_squared": r.get("r_squared"), "adj_r_squared": r.get("adj_r_squared"),
            "f_statistic": r.get("f_statistic"), "f_p_value": r.get("f_p_value"),
            "coefficients": [{"variable": c.get("variable"), "estimate": c.get("estimate"),
                              "p_value": c.get("p_value")} for c in r.get("coefficients", [])],
            "interpretation": r.get("interpretation"),
        }
    elif family == "correlation" and stat.get("correlation_result"):
        r = stat["correlation_result"]
        result_summary = {
            "statistic": f"{r.get('statistic_label')}={r.get('statistic')}",
            "p_value": r.get("p_value"), "verdict": r.get("verdict"),
            "strength": r.get("correlation_strength"), "interpretation": r.get("interpretation"),
        }
    elif family == "dimensionality" and stat.get("dimensionality_result"):
        r = stat["dimensionality_result"]
        result_summary = {
            "n_components_selected": r.get("n_components_selected"),
            "total_variance_explained": r.get("total_variance_explained"),
            "interpretation": r.get("interpretation"),
        }

    return json.dumps({
        "original_query": store.get("original_query"), "test_name": stat.get("test_name"),
        "test_family": family, "dependent_variable": meth.get("dependent_variable"),
        "independent_variables": meth.get("independent_variables", []),
        "n_rows": prof.get("n_rows"), "assumptions_passed": check.get("passed_count"),
        "assumptions_failed": check.get("failed_count"),
        "post_test_failures": crit.get("failed_count", 0), "result": result_summary,
    }, indent=2)


@tool
def build_and_render_report(interpretation: str) -> str:
    """
    Assembles the full report. Provide a 2-4 sentence plain English interpretation
    of what the results mean for the user's original question.
    """
    store = _report_store
    try:
        report = assemble_report(
            original_query=store.get("original_query", ""),
            profiler_output=store.get("profiler_output", {}),
            preprocessor_output=store.get("preprocessor_output", {}),
            methodologist_output=store.get("methodologist_output", {}),
            checker_output=store.get("checker_output", {}),
            statistician_output=store.get("statistician_output", {}),
            rectification_output=store.get("rectification_output"),
            critic_output=store.get("critic_output"),
            interpretation=interpretation,
        )
        _report_store["report_output"] = report
        return report.markdown_report
    except Exception as e:
        return f"ERROR: Could not assemble report — {str(e)}"


@tool
def generate_docx_report() -> str:
    """
    Generates a comprehensive .docx Word document with ALL pipeline data.
    Call this AFTER build_and_render_report.
    """
    report: FinalReportOutput | None = _report_store.get("report_output")
    if report is None:
        return "ERROR: No report assembled. Call build_and_render_report first."

    # ── Pull all raw pipeline outputs ──
    prof  = _report_store.get("profiler_output", {})
    prep  = _report_store.get("preprocessor_output", {})
    meth  = _report_store.get("methodologist_output", {})
    check = _report_store.get("checker_output", {})
    stat  = _report_store.get("statistician_output", {})
    crit  = _report_store.get("critic_output") or {}
    rect  = _report_store.get("rectification_output") or {}

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # ── Page margins ──
        for sec in doc.sections:
            sec.top_margin    = Inches(1)
            sec.bottom_margin = Inches(1)
            sec.left_margin   = Inches(1)
            sec.right_margin  = Inches(1)

        # ══════════════════════════════════════
        # HELPER FUNCTIONS
        # ══════════════════════════════════════

        def add_h1(text: str):
            h = doc.add_heading(text, level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
                run.font.size = Pt(18)
            return h

        def add_h2(text: str):
            h = doc.add_heading(text, level=2)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
                run.font.size = Pt(14)
            return h

        def add_h3(text: str):
            h = doc.add_heading(text, level=3)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1F, 0x50, 0x80)
                run.font.size = Pt(12)
            return h

        def add_kv(label: str, value):
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run("—" if value is None or value == "" else str(value))
            return p

        def add_divider():
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "C0C0C0")
            pBdr.append(bottom)
            pPr.append(pBdr)

        def fmt(val, decimals: int = 4) -> str:
            if val is None:
                return "—"
            try:
                f = float(val)
                return str(round(f, decimals))
            except Exception:
                return str(val)

        def add_table(headers: list, rows: list, col_widths: list = None):
            """
            Adds a formatted table with a blue header row and alternating row shading.
            col_widths: list of Inches values per column.
            """
            if not rows:
                doc.add_paragraph("(No data)")
                return

            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"

            # ── Header row ──
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                run = hdr_cells[i].paragraphs[0].runs[0]
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _shade_cell(hdr_cells[i], "2E75B6")

            # ── Data rows ──
            for r_idx, row_data in enumerate(rows):
                fill = "EBF3FB" if r_idx % 2 == 0 else "FFFFFF"
                cells = table.rows[r_idx + 1].cells
                for c_idx, val in enumerate(row_data):
                    cells[c_idx].text = "—" if (val is None or val == "") else str(val)
                    cells[c_idx].paragraphs[0].runs[0].font.size = Pt(10)
                    _shade_cell(cells[c_idx], fill)

            # ── Column widths ──
            if col_widths:
                for row in table.rows:
                    for i, cell in enumerate(row.cells):
                        if i < len(col_widths):
                            cell.width = Inches(col_widths[i])

            doc.add_paragraph()
            return table

        def _shade_cell(cell, hex_color: str):
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  hex_color)
            tcPr.append(shd)

        def status_icon(status: str) -> str:
            return {
                "passed":  "✅ Passed",
                "failed":  "❌ Failed",
                "warning": "⚠️ Warning",
                "manual":  "🔵 Manual",
            }.get(str(status).lower(), str(status))

        def sig_label(verdict: str) -> str:
            return {
                "significant":     "✅ Significant",
                "not_significant": "❌ Not Significant",
                "borderline":      "⚠️ Borderline",
            }.get(str(verdict).lower(), str(verdict))

        # ══════════════════════════════════════
        # COVER — Title + overview
        # ══════════════════════════════════════
        add_h1(report.title)

        p = doc.add_paragraph()
        p.add_run("Query: ").bold = True
        p.add_run(report.original_query)

        add_kv("Test Performed",    stat.get("test_name", "—"))
        add_kv("Total Observations", stat.get("n_observations") or prof.get("n_rows", "—"))
        add_kv("Dependent Variable",  meth.get("dependent_variable", "—"))

        ind = meth.get("independent_variables", [])
        if ind:
            add_kv("Independent Variables", ", ".join(ind))

        grp = meth.get("grouping_variable")
        if grp:
            add_kv("Grouping Variable", grp)

        add_divider()
        doc.add_paragraph()

        # ══════════════════════════════════════
        # SECTION 1 — Dataset Overview
        # ══════════════════════════════════════
        add_h2("1. Dataset Overview")
        add_kv("Total Rows",          prof.get("n_rows", "—"))
        add_kv("Total Columns",       prof.get("n_cols", "—"))
        add_kv("Total Missing Cells", prof.get("total_missing_cells", 0))
        add_kv("Overall Missing %",   f"{round(prof.get('total_missing_pct', 0) * 100, 2)}%")

        rows_dropped = prep.get("rows_dropped_total", 0)
        final_shape  = prep.get("final_shape", [prof.get("n_rows"), prof.get("n_cols")])
        if rows_dropped:
            add_kv("Rows After Cleaning", f"{final_shape[0]} ({rows_dropped} dropped)")
            changes = prep.get("changes_summary", [])
            if changes:
                add_h3("Cleaning Steps Applied")
                for ch in changes:
                    doc.add_paragraph(ch, style="List Bullet")
        else:
            doc.add_paragraph("✅ No data cleaning was required — dataset was already clean.")
        doc.add_paragraph()

        # ── Continuous Variables Table ──
        cont_cols = prof.get("continuous_columns", [])
        if cont_cols:
            add_h3("Continuous Variables — Descriptive Statistics")
            headers = ["Column", "Mean", "Median", "Std Dev", "Min", "Max",
                       "Skewness", "Kurtosis", "Missing %", "Outliers"]
            rows = [[
                c.get("column", ""),
                fmt(c.get("mean")),
                fmt(c.get("median")),
                fmt(c.get("std")),
                fmt(c.get("min")),
                fmt(c.get("max")),
                f"{fmt(c.get('skewness'))} ({c.get('skewness_interpretation', '') or ''})",
                fmt(c.get("kurtosis")),
                f"{round(c.get('missing_pct', 0) * 100, 1)}%",
                f"{c.get('anomaly_count', 0)} ({round(c.get('anomaly_pct', 0) * 100, 1)}%)",
            ] for c in cont_cols]
            add_table(headers, rows,
                      [1.3, 0.75, 0.75, 0.75, 0.65, 0.65, 1.4, 0.75, 0.75, 0.8])

            # Quartiles
            add_h3("Continuous Variables — Quartiles & Confidence Intervals")
            q_headers = ["Column", "Q1", "Q3", "IQR", "Range", "95% CI Lower", "95% CI Upper"]
            q_rows = []
            for c in cont_cols:
                ci = c.get("confidence_interval_95") or [None, None]
                q_rows.append([
                    c.get("column", ""),
                    fmt(c.get("q1")),
                    fmt(c.get("q3")),
                    fmt(c.get("iqr")),
                    fmt(c.get("range")),
                    fmt(ci[0] if len(ci) > 0 else None),
                    fmt(ci[1] if len(ci) > 1 else None),
                ])
            add_table(q_headers, q_rows, [1.5, 0.9, 0.9, 0.9, 0.9, 1.2, 1.2])

        # ── Categorical Variables Table ──
        cat_cols = prof.get("categorical_columns", [])
        if cat_cols:
            add_h3("Categorical Variables — Summary")
            cat_headers = ["Column", "Unique Values", "Mode", "Mode Freq %",
                           "Class Imbalance", "Missing %"]
            cat_rows = [[
                c.get("column", ""),
                c.get("cardinality", "—"),
                c.get("mode", "—"),
                f"{round((c.get('mode_frequency') or 0) * 100, 1)}%",
                "⚠️ Yes" if c.get("class_imbalance_flag") else "✅ No",
                f"{round(c.get('missing_pct', 0) * 100, 1)}%",
            ] for c in cat_cols]
            add_table(cat_headers, cat_rows, [1.5, 1.0, 1.2, 1.0, 1.2, 1.0])

            # Value distributions
            for c in cat_cols:
                vc = c.get("value_counts", {})
                if vc:
                    add_h3(f"Value Distribution — {c.get('column')}")
                    total = sum(vc.values()) or 1
                    vc_rows = [
                        [k, str(v), f"{round(v / total * 100, 1)}%"]
                        for k, v in sorted(vc.items(), key=lambda x: -x[1])
                    ]
                    add_table(["Value", "Count", "Percentage"], vc_rows, [3.0, 1.5, 1.5])

        # ── Data quality warnings ──
        prof_warnings = prof.get("warnings", [])
        if prof_warnings:
            add_h3("Data Quality Warnings")
            for w in prof_warnings:
                doc.add_paragraph(w, style="List Bullet")

        add_divider()
        doc.add_paragraph()

        # ══════════════════════════════════════
        # SECTION 2 — Test Selection
        # ══════════════════════════════════════
        add_h2("2. Statistical Test Selection")
        add_kv("Selected Test",   meth.get("selected_test", "—"))
        add_kv("Selection Mode",  meth.get("selection_mode", "—"))
        add_kv("Reasoning",       meth.get("reasoning", "—"))

        override = meth.get("override_reason")
        if override:
            p = doc.add_paragraph()
            p.add_run("⚠️ Override Note: ").bold = True
            p.add_run(override)

        mismatch = meth.get("mismatch_warning")
        if mismatch:
            p = doc.add_paragraph()
            p.add_run("⚠️ Mismatch Warning: ").bold = True
            p.add_run(mismatch)

        add_divider()
        doc.add_paragraph()

        # ══════════════════════════════════════
        # SECTION 3 — Pre-Test Assumption Checks
        # ══════════════════════════════════════
        add_h2("3. Pre-Test Assumption Checks")
        add_kv("Total Assumptions",    check.get("total_assumptions", "—"))
        add_kv("Passed",               check.get("passed_count", "—"))
        add_kv("Failed",               check.get("failed_count", "—"))
        add_kv("Warnings",             check.get("warning_count", "—"))
        add_kv("Manual Confirmations", check.get("manual_count", "—"))
        doc.add_paragraph()

        assumption_results = check.get("results", [])
        if assumption_results:
            a_headers = ["Assumption", "Status", "Method", "Statistic", "p-value", "Finding"]
            a_rows = [[
                r.get("name", ""),
                status_icon(r.get("status", "")),
                r.get("test_used") or "—",
                fmt(r.get("statistic")) if r.get("statistic") is not None else "—",
                fmt(r.get("p_value"))   if r.get("p_value")   is not None else "—",
                r.get("plain_reason", ""),
            ] for r in assumption_results]
            add_table(a_headers, a_rows, [1.3, 1.0, 1.2, 0.8, 0.7, 2.4])

        # ── Rectifications ──
        if rect and rect.get("chosen_solution"):
            add_h3("Rectifications Applied")
            sol = rect.get("chosen_solution", {})
            add_kv("Solution", sol.get("description", "—"))
            transforms = rect.get("applied_transforms", [])
            if transforms:
                t_rows = [[
                    t.get("transform_type", ""),
                    t.get("column", ""),
                    t.get("description", ""),
                ] for t in transforms]
                add_table(["Transform", "Column", "Description"], t_rows, [1.8, 1.5, 3.2])

        add_divider()
        doc.add_paragraph()

        # ══════════════════════════════════════
        # SECTION 4 — Statistical Test Results
        # ══════════════════════════════════════
        add_h2("4. Statistical Test Results")
        add_kv("Test",               stat.get("test_name", "—"))
        add_kv("Observations Used",  stat.get("n_observations", "—"))
        if stat.get("correction_applied"):
            add_kv("Correction Applied", stat["correction_applied"])
        doc.add_paragraph()

        family = stat.get("test_family", "")

        # ── REGRESSION ──
        if family == "regression" and stat.get("regression_result"):
            r = stat["regression_result"]

            add_h3("Model Fit Statistics")
            fit_rows = [
                ["R²",              fmt(r.get("r_squared"))],
                ["Adjusted R²",     fmt(r.get("adj_r_squared"))],
                ["F-statistic",     fmt(r.get("f_statistic"))],
                ["F p-value",       fmt(r.get("f_p_value"))],
                ["AIC",             fmt(r.get("aic"))],
                ["BIC",             fmt(r.get("bic"))],
                ["RMSE",            fmt(r.get("rmse"))],
            ]
            add_table(["Metric", "Value"], fit_rows, [3.5, 3.0])

            if r.get("coefficients"):
                add_h3("Regression Coefficients")
                coef_rows = []
                for c in r["coefficients"]:
                    pv  = c.get("p_value")
                    sig = ("***" if pv is not None and pv < 0.001 else
                           "**"  if pv is not None and pv < 0.01  else
                           "*"   if pv is not None and pv < 0.05  else "")
                    coef_rows.append([
                        c.get("variable", ""),
                        fmt(c.get("estimate")),
                        fmt(c.get("std_error")),
                        fmt(c.get("t_statistic")),
                        fmt(pv),
                        f"{fmt(c.get('ci_lower'))} – {fmt(c.get('ci_upper'))}",
                        sig,
                    ])
                add_table(
                    ["Variable", "Estimate", "Std Error", "t-stat", "p-value", "95% CI", "Sig."],
                    coef_rows, [1.8, 0.9, 0.9, 0.8, 0.8, 1.5, 0.5]
                )
                p = doc.add_paragraph()
                p.add_run("Significance codes: ").bold = True
                p.add_run("*** p < 0.001   ** p < 0.01   * p < 0.05")

        # ── INFERENCE ──
        elif family == "inference" and stat.get("inference_result"):
            r = stat["inference_result"]

            add_h3("Test Statistics")
            stat_rows = [
                [f"{r.get('statistic_label', 'stat')}-statistic", fmt(r.get("statistic"))],
                ["p-value",           fmt(r.get("p_value"))],
                ["Alpha (α)",         fmt(r.get("alpha", 0.05))],
                ["Verdict",           sig_label(str(r.get("verdict", "")))],
                ["Degrees of Freedom", fmt(r.get("df"))],
                ["Effect Size",
                 f"{r.get('effect_size_label') or ''} = {fmt(r.get('effect_size'))}"],
            ]
            add_table(["Metric", "Value"], stat_rows, [3.5, 3.0])

            gs = r.get("group_stats", {})
            if gs:
                add_h3("Group Descriptive Statistics")
                gs_rows = [[
                    grp,
                    fmt(gs[grp].get("mean")),
                    fmt(gs[grp].get("std")),
                    str(gs[grp].get("n", "—")),
                ] for grp in gs]
                add_table(["Group", "Mean", "Std Dev", "N"], gs_rows,
                          [2.0, 1.5, 1.5, 1.5])

        # ── CORRELATION ──
        elif family == "correlation" and stat.get("correlation_result"):
            r = stat["correlation_result"]

            add_h3("Correlation Statistics")
            corr_rows = [
                [f"{r.get('statistic_label', 'r')}",
                 fmt(r.get("statistic"))],
                ["p-value",    fmt(r.get("p_value"))],
                ["Verdict",    sig_label(str(r.get("verdict", "")))],
                ["Strength",   (r.get("correlation_strength") or "—").capitalize()],
                ["95% CI",
                 f"{fmt(r.get('ci_lower'))} – {fmt(r.get('ci_upper'))}"],
            ]
            add_table(["Metric", "Value"], corr_rows, [3.5, 3.0])

        # ── PCA ──
        elif family == "dimensionality" and stat.get("dimensionality_result"):
            r = stat["dimensionality_result"]

            add_h3("PCA Summary")
            add_table(["Metric", "Value"], [
                ["Total Input Variables",      str(r.get("n_components_total", "—"))],
                ["Components Selected (≥80%)", str(r.get("n_components_selected", "—"))],
                ["Total Variance Explained",   f"{fmt(r.get('total_variance_explained'))}%"],
            ], [3.5, 3.0])

            components = r.get("components", [])
            if components:
                add_h3("Component Variance Table")
                comp_rows = [[
                    f"PC{c.get('component_number')}",
                    f"{fmt(c.get('explained_variance_pct'))}%",
                    f"{fmt(c.get('cumulative_variance_pct'))}%",
                    fmt(c.get("explained_variance")),
                ] for c in components]
                add_table(["Component", "Variance %", "Cumulative %", "Eigenvalue"],
                          comp_rows, [1.5, 1.5, 1.5, 1.5])

                all_vars = list(components[0].get("loadings", {}).keys())
                if all_vars:
                    add_h3("Component Loadings")
                    load_headers = ["Variable"] + [f"PC{c.get('component_number')}"
                                                   for c in components]
                    load_rows = [[v] + [fmt(c.get("loadings", {}).get(v))
                                        for c in components] for v in all_vars]
                    col_w = [2.0] + [max(0.6, 6.5 / len(components))] * len(components)
                    add_table(load_headers, load_rows, col_w)

        # ── Verdict paragraph ──
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Verdict: ").bold = True
        p.add_run(report.verdict)

        add_divider()
        doc.add_paragraph()

        # ══════════════════════════════════════
        # SECTION 5 — Post-Test Model Checks
        # ══════════════════════════════════════
        if crit.get("checks_applicable", False):
            add_h2("5. Post-Test Model Checks")
            add_kv("Total Checks", crit.get("total_checks", "—"))
            add_kv("Passed",       crit.get("passed_count", "—"))
            add_kv("Failed",       crit.get("failed_count", "—"))
            add_kv("Warnings",     crit.get("warning_count", "—"))
            doc.add_paragraph()

            critic_results = crit.get("results", [])
            if critic_results:
                c_headers = ["Check", "Status", "Method", "Statistic", "p-value", "Finding"]
                c_rows = [[
                    r.get("name", ""),
                    status_icon(r.get("status", "")),
                    r.get("test_used") or "—",
                    fmt(r.get("statistic")) if r.get("statistic") is not None else "—",
                    fmt(r.get("p_value"))   if r.get("p_value")   is not None else "—",
                    r.get("plain_reason", ""),
                ] for r in critic_results]
                add_table(c_headers, c_rows, [1.3, 1.0, 1.2, 0.8, 0.7, 2.4])

            add_divider()
            doc.add_paragraph()

        # ══════════════════════════════════════
        # SECTION 6 — Interpretation
        # ══════════════════════════════════════
        sec = 6 if crit.get("checks_applicable") else 5
        add_h2(f"{sec}. Interpretation")
        doc.add_paragraph(report.interpretation)

        if report.effect_size:
            p = doc.add_paragraph()
            p.add_run("Effect Size: ").bold = True
            p.add_run(report.effect_size)

        add_divider()
        doc.add_paragraph()

        # ══════════════════════════════════════
        # SECTION 7 — Caveats
        # ══════════════════════════════════════
        if report.caveats:
            sec += 1
            add_h2(f"{sec}. Caveats & Limitations")
            for c in report.caveats:
                doc.add_paragraph(c, style="List Bullet")
            doc.add_paragraph()
            add_divider()
            doc.add_paragraph()

        # ── Footer ──
        p = doc.add_paragraph("Generated by Aristostat — The Logical Inference Engine")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.font.size = Pt(9)

        # ── Save ──
        project_dir = os.path.dirname(os.path.abspath(__file__))
        output_dir  = os.path.join(project_dir, "..", "Output")
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, "aristostat_report.docx")
        doc.save(output_path)

        report.docx_path      = output_path
        report.docx_generated = True
        _report_store["report_output"] = report

        return f"SUCCESS: Report saved to {output_path}"

    except Exception as e:
        import traceback
        return f"ERROR: Could not generate docx — {str(e)}"


# ─────────────────────────────────────────────
# EXPORTED TOOL LIST
# ─────────────────────────────────────────────

FINAL_REPORT_TOOLS = [
    get_pipeline_summary,
    build_and_render_report,
    generate_docx_report,
]